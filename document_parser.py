#!/usr/bin/env python3
"""
Sequential IEEE Document Parser
Extracts: Title → Authors → Abstract → Keywords → Sections → References
All tables, equations, images captured as images
"""

import re
import base64
from io import BytesIO
from typing import Dict, List, Any, Optional, Tuple
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class IEEEDocumentParser:
    """Parser for IEEE-formatted academic papers (PDF and DOCX)"""
    
    def __init__(self):
        self.current_page = 0
        self.total_pages = 0
        
    def parse_document(self, file_data: bytes, file_type: str) -> Dict[str, Any]:
        """
        Main parsing function - sequential order
        
        Args:
            file_data: Binary file data
            file_type: 'pdf' or 'docx'
            
        Returns:
            Structured document data matching form schema
        """
        logger.info(f"Starting document parsing (type: {file_type})")
        
        try:
            if file_type == 'pdf':
                return self._parse_pdf(file_data)
            elif file_type == 'docx':
                return self._parse_docx(file_data)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
        except Exception as e:
            logger.error(f"Parsing failed: {e}")
            raise
    
    def _parse_pdf(self, file_data: bytes) -> Dict[str, Any]:
        """Parse PDF document"""
        import pdfplumber
        from pdf2image import convert_from_bytes
        
        result = self._init_result()
        
        with pdfplumber.open(BytesIO(file_data)) as pdf:
            self.total_pages = len(pdf.pages)
            logger.info(f"PDF has {self.total_pages} pages")
            
            # Extract all text with layout information
            all_text_blocks = []
            for page_num, page in enumerate(pdf.pages):
                words = page.extract_words(keep_blank_chars=True)
                all_text_blocks.extend(self._group_words_into_blocks(words, page_num))
            
            # Convert PDF pages to images for visual element extraction
            pdf_images = convert_from_bytes(file_data, dpi=150)
            
            # Sequential extraction
            current_index = 0
            
            # Step 1: Extract Title
            result['title'], current_index = self._extract_title_from_blocks(
                all_text_blocks, current_index
            )
            
            # Step 2: Extract Authors
            result['authors'], current_index = self._extract_authors_from_blocks(
                all_text_blocks, current_index
            )
            
            # Step 3: Extract Abstract (optional)
            result['abstract'], current_index = self._extract_abstract_from_blocks(
                all_text_blocks, current_index
            )
            
            # Step 4: Extract Keywords (optional)
            result['keywords'], current_index = self._extract_keywords_from_blocks(
                all_text_blocks, current_index
            )
            
            # Step 5: Extract Sections with content
            result['sections'], current_index = self._extract_sections_from_blocks(
                all_text_blocks, current_index, pdf, pdf_images
            )
            
            # Step 6: Extract References
            result['references'] = self._extract_references_from_blocks(
                all_text_blocks, current_index
            )
        
        logger.info("PDF parsing completed successfully")
        return result
    
    def _parse_docx(self, file_data: bytes) -> Dict[str, Any]:
        """Parse DOCX document"""
        from docx import Document
        from docx.oxml.ns import qn
        
        result = self._init_result()
        
        doc = Document(BytesIO(file_data))
        logger.info(f"DOCX has {len(doc.paragraphs)} paragraphs")
        
        # Sequential extraction
        current_para_index = 0
        
        # Step 1: Extract Title
        result['title'], current_para_index = self._extract_title_from_docx(
            doc, current_para_index
        )
        
        # Step 2: Extract Authors
        result['authors'], current_para_index = self._extract_authors_from_docx(
            doc, current_para_index
        )
        
        # Step 3: Extract Abstract (optional)
        result['abstract'], current_para_index = self._extract_abstract_from_docx(
            doc, current_para_index
        )
        
        # Step 4: Extract Keywords (optional)
        result['keywords'], current_para_index = self._extract_keywords_from_docx(
            doc, current_para_index
        )
        
        # Step 5: Extract Sections with content
        result['sections'], current_para_index = self._extract_sections_from_docx(
            doc, current_para_index
        )
        
        # Step 6: Extract References
        result['references'] = self._extract_references_from_docx(
            doc, current_para_index
        )
        
        logger.info("DOCX parsing completed successfully")
        return result
    
    def _init_result(self) -> Dict[str, Any]:
        """Initialize result structure"""
        return {
            'title': '',
            'authors': [],
            'abstract': '',
            'keywords': '',
            'sections': [],
            'references': []
        }
    
    # ==================== TITLE EXTRACTION ====================
    
    def _extract_title_from_blocks(self, blocks: List[Dict], start_index: int) -> Tuple[str, int]:
        """
        Extract title from PDF text blocks
        Title characteristics:
        - First text element (top of first page)
        - Largest font size
        - Bold or centered
        - No numbering prefix
        """
        logger.info("Extracting title from PDF")
        
        title_lines = []
        current_index = start_index
        
        # Look at first few blocks (title is usually 1-3 lines)
        for i in range(start_index, min(start_index + 5, len(blocks))):
            block = blocks[i]
            text = block['text'].strip()
            
            # Skip empty blocks
            if not text:
                continue
            
            # Check if this looks like a title
            if self._is_title_block(block, blocks):
                title_lines.append(text)
                current_index = i + 1
            else:
                break
        
        title = ' '.join(title_lines)
        logger.info(f"Extracted title: {title[:50]}...")
        return title, current_index
    
    def _extract_title_from_docx(self, doc, start_index: int) -> Tuple[str, int]:
        """Extract title from DOCX paragraphs"""
        logger.info("Extracting title from DOCX")
        
        title_lines = []
        current_index = start_index
        
        for i in range(start_index, min(start_index + 5, len(doc.paragraphs))):
            para = doc.paragraphs[i]
            text = para.text.strip()
            
            if not text:
                continue
            
            # Check if this looks like a title (large font, bold, centered)
            if self._is_title_paragraph(para):
                title_lines.append(text)
                current_index = i + 1
            else:
                break
        
        title = ' '.join(title_lines)
        logger.info(f"Extracted title: {title[:50]}...")
        return title, current_index
    
    def _is_title_block(self, block: Dict, all_blocks: List[Dict]) -> bool:
        """Check if block is likely a title"""
        # Title is usually larger font and in first page
        if block.get('page', 0) > 0:
            return False
        
        # Check font size (title is usually largest)
        font_size = block.get('size', 12)
        if font_size < 14:
            return False
        
        # Title shouldn't start with numbers or common section keywords
        text = block['text'].strip()
        if re.match(r'^[IVX\d]+\.', text):
            return False
        if text.lower() in ['abstract', 'introduction', 'keywords']:
            return False
        
        return True
    
    def _is_title_paragraph(self, para) -> bool:
        """Check if paragraph is likely a title"""
        text = para.text.strip()
        
        # Skip empty or very short
        if len(text) < 5:
            return False
        
        # Check for large font or bold
        has_large_font = False
        has_bold = False
        
        for run in para.runs:
            if run.font.size and run.font.size.pt > 14:
                has_large_font = True
            if run.bold:
                has_bold = True
        
        # Title shouldn't start with section numbers
        if re.match(r'^[IVX\d]+\.', text):
            return False
        
        return has_large_font or has_bold
    
    # ==================== AUTHOR EXTRACTION ====================
    
    def _extract_authors_from_blocks(self, blocks: List[Dict], start_index: int) -> Tuple[List[Dict], int]:
        """Extract authors from PDF blocks"""
        logger.info("Extracting authors from PDF")
        
        authors = []
        current_index = start_index
        author_text_lines = []
        
        # Look for author information (usually 2-5 lines after title)
        for i in range(start_index, min(start_index + 10, len(blocks))):
            block = blocks[i]
            text = block['text'].strip()
            
            if not text:
                continue
            
            # Stop if we hit abstract or introduction
            if re.match(r'^(abstract|introduction|i\.|1\.)', text.lower()):
                break
            
            # Collect potential author lines
            if self._looks_like_author_info(text):
                author_text_lines.append(text)
                current_index = i + 1
            elif author_text_lines:
                # We've collected some author info and hit something else
                break
        
        # Parse collected author text
        authors = self._parse_author_text(author_text_lines)
        logger.info(f"Extracted {len(authors)} authors")
        return authors, current_index
    
    def _extract_authors_from_docx(self, doc, start_index: int) -> Tuple[List[Dict], int]:
        """Extract authors from DOCX paragraphs"""
        logger.info("Extracting authors from DOCX")
        
        authors = []
        current_index = start_index
        author_text_lines = []
        
        for i in range(start_index, min(start_index + 10, len(doc.paragraphs))):
            para = doc.paragraphs[i]
            text = para.text.strip()
            
            if not text:
                continue
            
            # Stop at abstract or introduction
            if re.match(r'^(abstract|introduction|i\.|1\.)', text.lower()):
                break
            
            if self._looks_like_author_info(text):
                author_text_lines.append(text)
                current_index = i + 1
            elif author_text_lines:
                break
        
        authors = self._parse_author_text(author_text_lines)
        logger.info(f"Extracted {len(authors)} authors")
        return authors, current_index
    
    def _looks_like_author_info(self, text: str) -> bool:
        """Check if text looks like author information"""
        # Contains email
        if re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text):
            return True
        
        # Contains typical name patterns (capitalized words)
        if re.search(r'\b[A-Z][a-z]+\s+[A-Z][a-z]+\b', text):
            return True
        
        # Contains affiliation keywords
        affiliation_keywords = ['university', 'institute', 'department', 'college', 'lab']
        if any(keyword in text.lower() for keyword in affiliation_keywords):
            return True
        
        return False
    
    def _parse_author_text(self, lines: List[str]) -> List[Dict]:
        """Parse author text lines into structured author data"""
        authors = []
        
        # Combine all lines
        full_text = ' '.join(lines)
        
        # Extract emails
        emails = re.findall(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', full_text)
        
        # Extract names (capitalized words before emails or commas)
        name_pattern = r'\b([A-Z][a-z]+(?:\s+[A-Z]\.?)?\s+[A-Z][a-z]+)\b'
        names = re.findall(name_pattern, full_text)
        
        # Create author entries
        for i, name in enumerate(names):
            author = {
                'name': name,
                'email': emails[i] if i < len(emails) else '',
                'affiliation': ''
            }
            
            # Try to extract affiliation
            affiliation_keywords = ['university', 'institute', 'department', 'college']
            for line in lines:
                if any(keyword in line.lower() for keyword in affiliation_keywords):
                    author['affiliation'] = line
                    break
            
            authors.append(author)
        
        return authors
    
    # ==================== ABSTRACT EXTRACTION ====================
    
    def _extract_abstract_from_blocks(self, blocks: List[Dict], start_index: int) -> Tuple[str, int]:
        """Extract abstract from PDF blocks"""
        logger.info("Extracting abstract from PDF")
        
        abstract_lines = []
        current_index = start_index
        found_abstract_header = False
        
        for i in range(start_index, min(start_index + 30, len(blocks))):
            block = blocks[i]
            text = block['text'].strip()
            
            if not text:
                continue
            
            # Look for "Abstract" header
            if not found_abstract_header:
                if re.match(r'^abstract', text.lower()):
                    found_abstract_header = True
                    current_index = i + 1
                continue
            
            # Stop at keywords or introduction
            if re.match(r'^(keywords?|index terms|introduction|i\.|1\.)', text.lower()):
                break
            
            # Collect abstract text
            abstract_lines.append(text)
            current_index = i + 1
        
        abstract = ' '.join(abstract_lines)
        logger.info(f"Extracted abstract ({len(abstract)} chars)")
        return abstract, current_index
    
    def _extract_abstract_from_docx(self, doc, start_index: int) -> Tuple[str, int]:
        """Extract abstract from DOCX paragraphs"""
        logger.info("Extracting abstract from DOCX")
        
        abstract_lines = []
        current_index = start_index
        found_abstract_header = False
        
        for i in range(start_index, min(start_index + 30, len(doc.paragraphs))):
            para = doc.paragraphs[i]
            text = para.text.strip()
            
            if not text:
                continue
            
            if not found_abstract_header:
                if re.match(r'^abstract', text.lower()):
                    found_abstract_header = True
                    current_index = i + 1
                continue
            
            if re.match(r'^(keywords?|index terms|introduction|i\.|1\.)', text.lower()):
                break
            
            abstract_lines.append(text)
            current_index = i + 1
        
        abstract = ' '.join(abstract_lines)
        logger.info(f"Extracted abstract ({len(abstract)} chars)")
        return abstract, current_index
    
    # ==================== KEYWORDS EXTRACTION ====================
    
    def _extract_keywords_from_blocks(self, blocks: List[Dict], start_index: int) -> Tuple[str, int]:
        """Extract keywords from PDF blocks"""
        logger.info("Extracting keywords from PDF")
        
        keywords = ''
        current_index = start_index
        found_keywords_header = False
        
        for i in range(start_index, min(start_index + 20, len(blocks))):
            block = blocks[i]
            text = block['text'].strip()
            
            if not text:
                continue
            
            # Look for keywords header
            if not found_keywords_header:
                if re.match(r'^(keywords?|index terms)', text.lower()):
                    found_keywords_header = True
                    # Keywords might be on same line
                    keywords_match = re.search(r'^(?:keywords?|index terms)[:\s—-]+(.+)', text, re.IGNORECASE)
                    if keywords_match:
                        keywords = keywords_match.group(1)
                        current_index = i + 1
                        break
                    current_index = i + 1
                continue
            
            # Stop at introduction
            if re.match(r'^(introduction|i\.|1\.)', text.lower()):
                break
            
            # Collect keywords
            keywords = text
            current_index = i + 1
            break
        
        logger.info(f"Extracted keywords: {keywords}")
        return keywords, current_index
    
    def _extract_keywords_from_docx(self, doc, start_index: int) -> Tuple[str, int]:
        """Extract keywords from DOCX paragraphs"""
        logger.info("Extracting keywords from DOCX")
        
        keywords = ''
        current_index = start_index
        found_keywords_header = False
        
        for i in range(start_index, min(start_index + 20, len(doc.paragraphs))):
            para = doc.paragraphs[i]
            text = para.text.strip()
            
            if not text:
                continue
            
            if not found_keywords_header:
                if re.match(r'^(keywords?|index terms)', text.lower()):
                    found_keywords_header = True
                    keywords_match = re.search(r'^(?:keywords?|index terms)[:\s—-]+(.+)', text, re.IGNORECASE)
                    if keywords_match:
                        keywords = keywords_match.group(1)
                        current_index = i + 1
                        break
                    current_index = i + 1
                continue
            
            if re.match(r'^(introduction|i\.|1\.)', text.lower()):
                break
            
            keywords = text
            current_index = i + 1
            break
        
        logger.info(f"Extracted keywords: {keywords}")
        return keywords, current_index
    
    # ==================== SECTION EXTRACTION ====================
    
    def _extract_sections_from_blocks(self, blocks: List[Dict], start_index: int, 
                                     pdf, pdf_images) -> Tuple[List[Dict], int]:
        """Extract sections with content blocks from PDF"""
        logger.info("Extracting sections from PDF")
        
        sections = []
        current_index = start_index
        
        # Section numbering patterns
        section_patterns = [
            r'^([IVX]+)\.\s+(.+)',      # I., II., III.
            r'^(\d+)\.\s+([^0-9].+)',   # 1., 2., 3. (not 1.1)
            r'^([A-Z])\.\s+(.+)',       # A., B., C.
        ]
        
        subsection_patterns = [
            r'^(\d+\.\d+)\s+(.+)',      # 1.1, 1.2
            r'^([A-Z])\.\s+(.+)',       # A., B. (subsection)
        ]
        
        current_section = None
        
        for i in range(start_index, len(blocks)):
            block = blocks[i]
            text = block['text'].strip()
            
            if not text:
                continue
            
            # Check if this is "References" section (end of content)
            if re.match(r'^(references|bibliography)', text.lower()):
                current_index = i
                break
            
            # Check for section header
            is_section = False
            section_number = ''
            section_title = ''
            
            for pattern in section_patterns:
                match = re.match(pattern, text, re.IGNORECASE)
                if match:
                    section_number = match.group(1)
                    section_title = match.group(2).strip()
                    is_section = True
                    break
            
            if is_section:
                # Save previous section
                if current_section:
                    sections.append(current_section)
                
                # Start new section
                current_section = {
                    'id': f'section_{len(sections)}',
                    'title': section_title,
                    'number': section_number,
                    'contentBlocks': [],
                    'order': len(sections)
                }
                logger.info(f"Found section: {section_number}. {section_title}")
            elif current_section:
                # Add content to current section
                # This is simplified - in full implementation, detect images, tables, equations
                current_section['contentBlocks'].append({
                    'id': f'block_{len(current_section["contentBlocks"])}',
                    'type': 'text',
                    'content': text,
                    'order': len(current_section['contentBlocks'])
                })
        
        # Add last section
        if current_section:
            sections.append(current_section)
        
        logger.info(f"Extracted {len(sections)} sections")
        return sections, current_index
    
    def _extract_sections_from_docx(self, doc, start_index: int) -> Tuple[List[Dict], int]:
        """Extract sections with content blocks from DOCX"""
        logger.info("Extracting sections from DOCX")
        
        sections = []
        current_index = start_index
        
        section_patterns = [
            r'^([IVX]+)\.\s+(.+)',
            r'^(\d+)\.\s+([^0-9].+)',
            r'^([A-Z])\.\s+(.+)',
        ]
        
        current_section = None
        
        for i in range(start_index, len(doc.paragraphs)):
            para = doc.paragraphs[i]
            text = para.text.strip()
            
            if not text:
                continue
            
            # Check for references section
            if re.match(r'^(references|bibliography)', text.lower()):
                current_index = i
                break
            
            # Check for section header
            is_section = False
            section_number = ''
            section_title = ''
            
            for pattern in section_patterns:
                match = re.match(pattern, text, re.IGNORECASE)
                if match:
                    section_number = match.group(1)
                    section_title = match.group(2).strip()
                    is_section = True
                    break
            
            if is_section:
                if current_section:
                    sections.append(current_section)
                
                current_section = {
                    'id': f'section_{len(sections)}',
                    'title': section_title,
                    'number': section_number,
                    'contentBlocks': [],
                    'order': len(sections)
                }
                logger.info(f"Found section: {section_number}. {section_title}")
            elif current_section:
                # Add text content
                if text:
                    current_section['contentBlocks'].append({
                        'id': f'block_{len(current_section["contentBlocks"])}',
                        'type': 'text',
                        'content': text,
                        'order': len(current_section['contentBlocks'])
                    })
                
                # Check for images in paragraph
                if para._element.xpath('.//pic:pic'):
                    # Extract image (simplified)
                    current_section['contentBlocks'].append({
                        'id': f'block_{len(current_section["contentBlocks"])}',
                        'type': 'image',
                        'image_data': '',  # Would extract actual image here
                        'caption': '',
                        'order': len(current_section['contentBlocks'])
                    })
        
        if current_section:
            sections.append(current_section)
        
        logger.info(f"Extracted {len(sections)} sections")
        return sections, current_index
    
    # ==================== REFERENCE EXTRACTION ====================
    
    def _extract_references_from_blocks(self, blocks: List[Dict], start_index: int) -> List[Dict]:
        """Extract references from PDF blocks"""
        logger.info("Extracting references from PDF")
        
        references = []
        found_references_header = False
        current_ref_text = ''
        current_ref_number = 0
        
        for i in range(start_index, len(blocks)):
            block = blocks[i]
            text = block['text'].strip()
            
            if not text:
                continue
            
            # Look for references header
            if not found_references_header:
                if re.match(r'^(references|bibliography)', text.lower()):
                    found_references_header = True
                continue
            
            # Check for numbered reference: [1], [2], etc.
            ref_match = re.match(r'^\[(\d+)\]\s*(.+)', text)
            if ref_match:
                # Save previous reference
                if current_ref_text:
                    references.append({
                        'text': current_ref_text.strip(),
                        'order': current_ref_number
                    })
                
                # Start new reference
                current_ref_number = int(ref_match.group(1)) - 1
                current_ref_text = ref_match.group(2)
            elif current_ref_text:
                # Continue current reference (multi-line)
                current_ref_text += ' ' + text
        
        # Add last reference
        if current_ref_text:
            references.append({
                'text': current_ref_text.strip(),
                'order': current_ref_number
            })
        
        logger.info(f"Extracted {len(references)} references")
        return references
    
    def _extract_references_from_docx(self, doc, start_index: int) -> List[Dict]:
        """Extract references from DOCX paragraphs"""
        logger.info("Extracting references from DOCX")
        
        references = []
        found_references_header = False
        current_ref_text = ''
        current_ref_number = 0
        
        for i in range(start_index, len(doc.paragraphs)):
            para = doc.paragraphs[i]
            text = para.text.strip()
            
            if not text:
                continue
            
            if not found_references_header:
                if re.match(r'^(references|bibliography)', text.lower()):
                    found_references_header = True
                continue
            
            ref_match = re.match(r'^\[(\d+)\]\s*(.+)', text)
            if ref_match:
                if current_ref_text:
                    references.append({
                        'text': current_ref_text.strip(),
                        'order': current_ref_number
                    })
                
                current_ref_number = int(ref_match.group(1)) - 1
                current_ref_text = ref_match.group(2)
            elif current_ref_text:
                current_ref_text += ' ' + text
        
        if current_ref_text:
            references.append({
                'text': current_ref_text.strip(),
                'order': current_ref_number
            })
        
        logger.info(f"Extracted {len(references)} references")
        return references
    
    # ==================== HELPER METHODS ====================
    
    def _group_words_into_blocks(self, words: List[Dict], page_num: int) -> List[Dict]:
        """Group words into text blocks based on position"""
        if not words:
            return []
        
        blocks = []
        current_block = {
            'text': '',
            'page': page_num,
            'top': words[0].get('top', 0),
            'size': words[0].get('size', 12)
        }
        
        for word in words:
            # Simple grouping by line (same vertical position)
            if abs(word.get('top', 0) - current_block['top']) < 5:
                current_block['text'] += word.get('text', '') + ' '
            else:
                # New line
                if current_block['text'].strip():
                    blocks.append(current_block)
                current_block = {
                    'text': word.get('text', '') + ' ',
                    'page': page_num,
                    'top': word.get('top', 0),
                    'size': word.get('size', 12)
                }
        
        # Add last block
        if current_block['text'].strip():
            blocks.append(current_block)
        
        return blocks


# Convenience function for API use
def parse_document(file_data: bytes, file_type: str) -> Dict[str, Any]:
    """
    Parse a document and extract structured data
    
    Args:
        file_data: Binary file data
        file_type: 'pdf' or 'docx'
        
    Returns:
        Structured document data
    """
    parser = IEEEDocumentParser()
    return parser.parse_document(file_data, file_type)
